import os
import json
from docx import Document
from pathlib import Path

# Configuration
PROJECT_ROOT = Path("/Users/joshuabhawanlall/Git Folder/Clinical_AI_Extraction_Pipeline")
DOCX_PATH = PROJECT_ROOT / "data/hackathon-mdt-outcome-proformas.docx"
OUTPUT_DIR = PROJECT_ROOT / "OCR-NER-Pipeline-v3/output/json_raw_v2"

def extract_content_to_json(docx_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"Loading document: {docx_path}")
    doc = Document(docx_path)
    
    # We want to extract both tables AND the headers/paragraphs between them
    # because MDT dates often live in paragraphs outside the tables.
    
    cases = []
    current_case = None
    
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            # It's a table
            table_index = len(cases)
            table = doc.tables[table_index] if table_index < len(doc.tables) else None
            if not table: continue
            
            rows_data = []
            for row_index, row in enumerate(table.rows):
                cells_text = []
                seen = set()
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen:
                        cells_text.append(text)
                        seen.add(text)
                rows_data.append({"row_index": row_index, "text": " | ".join(cells_text)})
            
            case_data = {
                "case_index": table_index,
                "table_rows": rows_data,
                "context_paragraphs": []
            }
            cases.append(case_data)
            current_case = case_data
            
        elif element.tag.endswith('p'):
            # It's a paragraph - might be MDT date header
            from docx.text.paragraph import Paragraph
            p = Paragraph(element, doc)
            text = p.text.strip()
            if text and current_case:
                current_case["context_paragraphs"].append(text)
            elif text and not current_case:
                # First case header
                pass

    for i, case in enumerate(cases):
        output_file = output_dir / f"case_{i:03d}_raw.json"
        with open(output_file, "w") as f:
            json.dump(case, f, indent=4)
        
    print(f"Finished. 50 cases saved to {output_dir}")

if __name__ == "__main__":
    extract_content_to_json(DOCX_PATH, OUTPUT_DIR)
