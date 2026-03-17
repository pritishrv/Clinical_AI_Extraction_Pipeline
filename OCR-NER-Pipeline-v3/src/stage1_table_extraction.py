import os
import json
from docx import Document
from pathlib import Path

# Configuration
DOCX_PATH = Path("data/hackathon-mdt-outcome-proformas.docx")
OUTPUT_DIR = Path("Git Folder/Clinical_AI_Extraction_Pipeline/OCR-NER-Pipeline-v3/output/json_raw")

def extract_tables_to_json(docx_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"Loading document: {docx_path}")
    doc = Document(docx_path)
    
    # Each case is in a separate table
    print(f"Found {len(doc.tables)} tables.")
    
    for i, table in enumerate(doc.tables):
        case_data = {
            "case_index": i,
            "rows": []
        }
        
        for row_index, row in enumerate(table.rows):
            # Deduplicate text in cells (merged cells repeat text)
            cells_text = []
            seen = set()
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen:
                    cells_text.append(text)
                    seen.add(text)
            
            row_text = " | ".join(cells_text)
            case_data["rows"].append({
                "row_index": row_index,
                "text": row_text
            })
        
        output_file = output_dir / f"case_{i:03d}_raw.json"
        with open(output_file, "w") as f:
            json.dump(case_data, f, indent=4)
        
        if i % 10 == 0:
            print(f"Processed {i} cases...")

    print(f"Finished. JSON files saved to {output_dir}")

if __name__ == "__main__":
    # Get the project root directory
    project_root = Path("/Users/joshuabhawanlall/Git Folder/Clinical_AI_Extraction_Pipeline")
    
    # Define absolute paths
    docx_path = project_root / "data" / "hackathon-mdt-outcome-proformas.docx"
    output_dir = project_root / "OCR-NER-Pipeline-v3" / "output" / "json_raw"
    
    extract_tables_to_json(docx_path, output_dir)
