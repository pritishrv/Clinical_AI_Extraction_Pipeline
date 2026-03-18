import os
import json
import re
from docx import Document
from pathlib import Path

# Configuration
PROJECT_ROOT = Path("/Users/joshuabhawanlall/Git Folder/Clinical_AI_Extraction_Pipeline")
DOCX_PATH = PROJECT_ROOT / "data/hackathon-mdt-outcome-proformas.docx"
OUTPUT_DIR = PROJECT_ROOT / "OCR-NER-Pipeline-v3/output/master_harvest"

def harvest_everything(doc):
    """Pulls every single line of text from the entire document."""
    all_lines = []
    
    # 1. Capture Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text: all_lines.append(text)
        
    # 2. Capture Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text: all_lines.append(row_text)
            
    return all_lines

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    print(f"Master Harvesting from {DOCX_PATH}...")
    doc = Document(DOCX_PATH)
    
    # Each 'Case' is a separate table in this doc
    for i, table in enumerate(doc.tables):
        print(f"Mining Case {i}...")
        
        # Capture the context: Case table + nearest paragraphs
        # (Paragraphs before the table usually contain the MDT Date)
        case_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            case_text.append(" | ".join(cells))
            
        output_file = OUTPUT_DIR / f"case_{i:03d}_full.json"
        with open(output_file, "w") as f:
            json.dump({
                "case_index": i,
                "text": " || ".join(case_text)
            }, f, indent=4)

    print(f"\nMaster Harvesting Complete.")

if __name__ == "__main__":
    main()
