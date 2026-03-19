import os
import subprocess
from docx import Document
from pathlib import Path
from tqdm import tqdm

# Configuration
DOCX_PATH = Path("data/hackathon-mdt-outcome-proformas.docx")
SPLIT_DIR = Path("v5_version_1/output/split_docs")
IMAGE_DIR = Path("v5_version_1/output/images")

def get_elements_in_order(doc):
    """Yields paragraphs and tables in the order they appear in the XML."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    for element in doc.element.body:
        if element.tag.endswith('p'):
            yield Paragraph(element, doc)
        elif element.tag.endswith('tbl'):
            yield Table(element, doc)

def split_docx_with_headers():
    """Splits docx into cases, preserving the header paragraph (date) above each table."""
    if not os.path.exists(SPLIT_DIR): os.makedirs(SPLIT_DIR)
    
    doc = Document(DOCX_PATH)
    elements = list(get_elements_in_order(doc))
    
    cases = []
    current_header = None
    
    # 1. Group Headers with their Tables
    for el in elements:
        if hasattr(el, 'text') and "Multidisciplinary Meeting" in el.text:
            current_header = el.text
        elif hasattr(el, 'rows'): # It's a table
            cases.append({
                "header": current_header,
                "table": el
            })
            current_header = None # Reset for next
            
    # 2. Save each case as a separate Docx
    for i, case in enumerate(tqdm(cases, desc="Splitting Docx with Headers")):
        new_doc = Document()
        
        # Add the Header (Date)
        if case["header"]:
            p = new_doc.add_paragraph(case["header"])
            p.runs[0].bold = True
            
        # Add the Table
        old_table = case["table"]
        new_table = new_doc.add_table(rows=len(old_table.rows), cols=len(old_table.columns))
        new_table.style = old_table.style
        
        for r_idx, row in enumerate(old_table.rows):
            for c_idx, cell in enumerate(row.cells):
                new_table.cell(r_idx, c_idx).text = cell.text
        
        new_doc.save(SPLIT_DIR / f"case_{i:03d}.docx")

def generate_high_res_images():
    """Uses macOS Quick Look (qlmanage) to generate a high-fidelity PNG."""
    if not os.path.exists(IMAGE_DIR): os.makedirs(IMAGE_DIR)
    
    split_files = sorted([f for f in os.listdir(SPLIT_DIR) if f.endswith(".docx")])
    for f in tqdm(split_files, desc="Rendering PNGs"):
        input_path = SPLIT_DIR / f
        subprocess.run([
            "qlmanage", "-t", "-s", "2000", 
            "-o", str(IMAGE_DIR), 
            str(input_path)
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        old_name = IMAGE_DIR / f"{f}.png"
        new_name = IMAGE_DIR / f.replace(".docx", ".png")
        if os.path.exists(old_name):
            os.rename(old_name, new_name)

if __name__ == "__main__":
    print("Step 1: Splitting Document (Capturing Dates)...")
    split_docx_with_headers()
    print("Step 2: Rendering High-Fidelity Images...")
    generate_high_res_images()
    print("Done! Dates are now visible at the top of the PNGs.")
