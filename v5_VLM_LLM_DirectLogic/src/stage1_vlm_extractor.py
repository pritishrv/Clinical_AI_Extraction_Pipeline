import os
import json
import re
import time
from docx import Document
from pathlib import Path
from tqdm import tqdm
import google.generativeai as genai
from config import GOOGLE_API_KEY

# Configure Gemini Pro
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro')

# Configuration
DOCX_PATH = Path("data/hackathon-mdt-outcome-proformas.docx")
OUTPUT_DIR = Path("v5-version_!1/output/json_raw")
LOG_DIR = Path("v5-version_!1/output/logs")

# Target Schema (88 Columns)
TARGET_SCHEMA = [
    "Demographics: DOB(a)", "Demographics: Initials(b)", "Demographics: MRN(c)", 
    "Demographics: NHS number(d)", "Demographics: Gender(e)", 
    "Demographics: Previous cancer (y, n) if yes, where(f)", 
    "Demographics: State site of previous cancer(f)", "Endoscopy: date(f)", 
    "Endosopy type(f)", "Endoscopy: Findings(f)", "Histology: Biopsy result(g)", 
    "Histology: Biopsy date(g)", "Histology: MMR status(g/h)", "Baseline MRI: date(h)", 
    "Baseline MRI: mrT(h)", "Baseline MRI: mrN(h)", "Baseline MRI: mrEMVI(h)", 
    "Baseline MRI: mrCRM(h)", "Baseline MRI: mrPSW(h)", "Baseline CT: Date(h)", 
    "Baseline CT: T(h)", "Baseline CT: N(h)", "Baseline CT: EMVI(h)", 
    "Baseline CT: M(h)", "Baseline CT: Incidental findings(h)", 
    "Baseline CT: Detail incidental finding(h)", "1st MDT: date(i)", 
    "1st MDT: Treatment approach(h)", "Chemotherapy: Treatment goals(curative, palliative)", 
    "Chemotherapy: Drugs", "Chemotherapy: Cycles", "Chemotherapy: Dates", 
    "Chemotherapy: Breaks", "Immunotherapy: Dates", "Immunotherapy", 
    "Radiotheapy: Total dose", "Radiotheapy: Boost", "Radiotherapy: Dates", 
    "Radiotheapy: Concomittant chemotherapy", "CEA: Date", "CEA: Value", 
    "CEA: DRE date", "CEA: DRE finding", "Surgery: Defunctioned?", 
    "Surgery: Date of surgery", "Surgery: Intent, pre-neoadjuvant therapy", 
    "2nd MRI: Date", "2nd MRI: Patient pathway status", "2nd MRI: mrT", 
    "2nd MRI: mrN", "2nd MRI: mrEMVI", "2nd MRI: mrCRM", "2nd MRI: mrPSW", 
    "2nd MRI: mrTRG score", "MDT (after 6 week: Date", "MDT (after 6 week: Decision", 
    "12 week MRI: Date", "12 week MRI: mrT", "12 week MRI: mrN", 
    "12 week MRI: mrEMVI", "12 week MRI: mrCRM", "12 week MRI: mrPSW", 
    "12 week MRI: mrTRG score", "Flex sig: Date", "Flex sig: Findings", 
    "MDT (after 12 week): Date", "MDT (after 12 week): Decision", 
    "Watch and wait: Entered watch + wait, date of MDT?", 
    "Watch and wait: Why did they enter wait(with what intent)", 
    "Watch and wait: Frequency?", "Watch and wait: Date of progression", 
    "Watch and wait: Site of progression", "Watch and wait: Date of death", 
    "MRI and flexisigmoidoscopy dates: Date entered watch and wait", 
    "MRI and flexisigmoidoscopy dates: Flexisigmoidoscopy date 1", 
    "MRI and flexisigmoidoscopy dates: Due date next 1", 
    "MRI and flexisigmoidoscopy dates: Flexisigmoidoscopy date 2", 
    "MRI and flexisigmoidoscopy dates: Due date next 2", 
    "MRI and flexisigmoidoscopy dates: Flexisigmoidoscopy date 3", 
    "MRI and flexisigmoidoscopy dates: Due date next 3", 
    "MRI and flexisigmoidoscopy dates: Flexisigmoidoscopy date 4", 
    "MRI and flexisigmoidoscopy dates: Due date next 4", 
    "MRI and flexisigmoidoscopy dates: MRI Date 1", 
    "MRI and flexisigmoidoscopy dates: Due date next 5", 
    "MRI and flexisigmoidoscopy dates: MRI Date 2", 
    "MRI and flexisigmoidoscopy dates: Due date next 6", 
    "MRI and flexisigmoidoscopy dates: MRI Date 3", 
    "MRI and flexisigmoidoscopy dates: Due date next 7"
]

SYSTEM_PROMPT = """
You are a Clinical Data Auditor. Your task is to extract patient journey data from a Multidisciplinary Team (MDT) proforma.

SAFETY RULES:
1. NO HALLUCINATIONS: If a data point is not explicitly stated, you MUST return null.
2. EVIDENCE ANCHORING: For every field you populate, you must provide the 'evidence'—the exact verbatim snippet from the text.
3. NO INFERENCE: Do not assume a date or stage unless it is written.
4. JSON ONLY: Return a JSON object where each key in the schema maps to an object: {"value": "extracted_data", "evidence": "verbatim_text_snippet"}.
"""

def extract_vlm(case_text):
    # Construct a schema where every key expects a {value, evidence} object
    evidence_schema = {key: {"value": "string or null", "evidence": "string or null"} for key in TARGET_SCHEMA}
    
    prompt = f"{SYSTEM_PROMPT}\n\nTARGET SCHEMA STRUCTURE:\n{json.dumps(evidence_schema)}\n\nCASE TEXT:\n{case_text}"
    
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                max_output_tokens=8192, # Increased for evidence
                temperature=0.0, # Zero temperature for maximum determinism
                response_mime_type="application/json",
            )
        )
        raw_json = json.loads(response.text)
        
        # --- ANTI-HALLUCINATION VERIFICATION ---
        verified_json = {}
        for key, data in raw_json.items():
            if not data or data.get("value") is None:
                verified_json[key] = None
                continue
                
            val = data.get("value")
            evidence = data.get("evidence", "")
            
            # 1. Check if evidence actually exists in the source text
            if evidence and str(evidence).lower() in case_text.lower():
                verified_json[key] = val
                # Optional: log the evidence for audit
                # verified_json[f"{key}_evidence"] = evidence
            else:
                # Discard value if evidence is missing or halluncinated
                verified_json[key] = None
                
        return verified_json
    except Exception as e:
        print(f"Error extracting case: {e}")
        return None

def main():
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    if not os.path.exists(LOG_DIR): os.makedirs(LOG_DIR)
    
    doc = Document(DOCX_PATH)
    print(f"Processing {len(doc.tables)} MDT cases...")
    
    for i, table in enumerate(tqdm(doc.tables)):
        # Convert table to high-fidelity text representation
        case_rows = []
        for row in table.rows:
            case_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        case_text = "\n".join(case_rows)
        
        # Call VLM
        result = extract_vlm(case_text)
        
        if result:
            # Anchor by NHS number for unique identification
            nhs = str(result.get("Demographics: NHS number(d)", f"UNKNOWN_{i}"))
            nhs_clean = re.sub(r'\D', '', nhs) if nhs else f"UNKNOWN_{i}"
            
            with open(OUTPUT_DIR / f"case_{i:03d}_{nhs_clean}.json", "w") as f:
                json.dump(result, f, indent=4)
        
        # Basic rate limiting for safety
        time.sleep(1)

if __name__ == "__main__":
    main()
