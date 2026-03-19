import os
import json
import re
from docx import Document
from pathlib import Path

# Configuration for v5 Gemini Implementation
PROJECT_ROOT = Path("/Users/joshuabhawanlall/Git Folder/Clinical_AI_Extraction_Pipeline")
DOCX_PATH = PROJECT_ROOT / "data/hackathon-mdt-outcome-proformas.docx"
OUTPUT_DIR = PROJECT_ROOT / "v5-Gemini-Implementation/output/journey_json"

def clean_nhs(text):
    match = re.search(r"(\d{10})", str(text))
    return match.group(1) if match else "unknown"

def harvest_v5_multimodal(doc):
    """
    Ingests tables AND paragraphs, grouping them by patient and MDT meeting date.
    """
    journey_store = {}
    current_nhs = "unknown"
    current_mdt_date = "unknown"
    
    # Process the document element by element to maintain order
    for element in doc.element.body:
        # Handle Paragraphs (usually containing MDT Date)
        if element.tag.endswith('p'):
            text = "".join(element.itertext()).strip()
            if not text: continue
            
            # Check for MDT Date Anchor
            date_match = re.search(r"(\d{1,2}/\d{1,2}/\d{4})\(i\)", text)
            if date_match:
                current_mdt_date = date_match.group(1)
                
        # Handle Tables (containing Demographics and Outcomes)
        elif element.tag.endswith('tbl'):
            # Find the actual table object
            # (Crude index-based lookup for hackathon speed)
            pass

    # REFINED APPROACH: Use doc.tables but capture 'nearest' para text
    for i, table in enumerate(doc.tables):
        # 1. Identity
        raw_demo = table.rows[1].cells[0].text
        nhs = clean_nhs(raw_demo)
        
        if nhs not in journey_store:
            journey_store[nhs] = {"demographics": {}, "events": []}
            
        # 2. Sequential Prose (Recursive Table Mine)
        all_prose = []
        for row in table.rows:
            # Deduplicate text in merged cells
            row_vals = list(dict.fromkeys([c.text.strip() for c in row.cells if c.text.strip()]))
            all_prose.append(" | ".join(row_vals))
            
        full_text = " || ".join(all_prose)
        
        # 3. MDT Date (Anchor)
        mdt_date = "unknown"
        # Search the table and surrounding for (i) marker
        date_m = re.search(r"(\d{1,2}/\d{1,2}/\d{4})\(i\)", full_text)
        if date_m: mdt_date = date_m.group(1)
        
        # 4. CAPTURE CANCER TARGET DATES (Cluster Awareness)
        target_dates = []
        if "DAY TARGET" in full_text:
            target_dates = re.findall(r"(\d{1,2}/\d{1,2}/\d{4})", full_text)

        event = {
            "case_index": i,
            "mdt_date": mdt_date,
            "raw_prose": full_text,
            "target_dates": target_dates
        }
        journey_store[nhs]["events"].append(event)
        
    return journey_store

def main():
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    print(f"v5 Gemini Implementation: Starting Multimodal Harvesting...")
    doc = Document(DOCX_PATH)
    journey_data = harvest_v5_multimodal(doc)
    
    with open(OUTPUT_DIR / "master_journey_v5.json", "w") as f:
        json.dump(journey_data, f, indent=4)
    print(f"Master Journey JSON updated with {len(journey_data)} patients.")

if __name__ == "__main__":
    main()
