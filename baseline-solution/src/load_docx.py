import os
from docx import Document

def load_cases(file_path):
    """
    Loads the DOCX and segments it into cases.
    Each case is currently identified by a table in the document.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    print(f"Loaded document with {len(doc.tables)} tables.")

    cases = []
    
    # Each table seems to correspond to one case
    # We also want to capture the MDT date which is in the paragraph before or near the table
    
    # Heuristic: Find the paragraph before each table to get the MDT date
    for i, table in enumerate(doc.tables):
        case_data = {
            'index': i,
            'table': table,
            'header_text': ""
        }
        
        # Try to find the date in nearby paragraphs
        # This is a bit tricky with python-docx as tables aren't directly in the paragraphs list
        # But we saw that paragraphs 1, 7, 13, 19 contain the date.
        # Let's just collect all text from the table for now as it contains most info.
        cases.append(case_data)
        
    return cases, doc

if __name__ == "__main__":
    docx_path = "../data/hackathon-mdt-outcome-proformas.docx"
    cases, doc = load_cases(docx_path)
    print(f"Successfully segmented {len(cases)} cases.")
    
    # Inspect first case table content
    first_table = cases[0]['table']
    for row in first_table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        # print(row_text)
