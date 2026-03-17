import os
import pandas as pd
from docx import Document

def inspect_docx(file_path):
    print(f"--- Inspecting DOCX: {file_path} ---")
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    doc = Document(file_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")

    # Print first 20 paragraphs to see structure
    print("\nFirst 20 paragraphs:")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"[{i}] {para.text[:100]}...")

    # Print first table structure if exists
    if doc.tables:
        print("\nFirst table structure:")
        table = doc.tables[0]
        for i, row in enumerate(table.rows[:5]):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"Row {i}: {cells}")

def inspect_xlsx(file_path):
    print(f"\n--- Inspecting XLSX: {file_path} ---")
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    xl = pd.ExcelFile(file_path)
    print(f"Sheet names: {xl.sheet_names}")

    for sheet in xl.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet)
        print(f"\nSheet: {sheet}")
        print(f"Columns: {list(df.columns)}")
        print(f"Shape: {df.shape}")
        print("First 3 rows:")
        print(df.head(3).to_string())

if __name__ == "__main__":
    docx_path = "../data/hackathon-mdt-outcome-proformas.docx"
    xlsx_path = "../data/hackathon-database-prototype.xlsx"

    inspect_docx(docx_path)
    inspect_xlsx(xlsx_path)
