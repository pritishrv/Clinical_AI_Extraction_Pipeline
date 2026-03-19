"""
Validation Script: Evaluation of V5 Pipeline
============================================
Step 1 - Combined Approach:
  Layer A: JSON-based broad coverage validation (Stage 2 mapping quality)
  Layer B: Independent docx re-parsing for critical fields (ground truth audit)

Usage:
  python validate_v5.py \
    --docx  data/hackathon-mdt-outcome-proformas.docx \
    --excel v5_VLM_LLM_DirectLogic/output/hackathon_output_final.xlsx \
    --jsons v5_VLM_LLM_DirectLogic/output/json_raw_claude/
"""

import os
import re
import json
import argparse
import warnings
from pathlib import Path
from difflib import SequenceMatcher

import pandas as pd
import docx

warnings.filterwarnings("ignore")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def normalise_date(val):
    """Normalise dates to DD/MM/YYYY regardless of input format."""
    if not val:
        return ""
    s = str(val).strip()
    # ISO format YYYY-MM-DD → DD/MM/YYYY
    m = re.match(r'^(\d{4})-(\d{2})-(\d{2})', s)
    if m:
        return f"{m.group(3)}/{m.group(2)}/{m.group(1)}"
    # Already DD/MM/YYYY
    m = re.match(r'^(\d{2})/(\d{2})/(\d{4})', s)
    if m:
        return s[:10]
    return s


def normalise(val):
    """Strip whitespace, lowercase, remove label markers like (a)(b)..."""
    if val is None:
        return ""
    s = str(val).strip()
    s = re.sub(r'\([a-zA-Z]\)', '', s)   # remove (a), (b), (d) etc.
    s = re.sub(r'\s+', ' ', s).strip()
    return s.lower()

def fuzzy_score(a, b):
    """Token-sort fuzzy similarity 0-100 using difflib."""
    a_tokens = " ".join(sorted(normalise(a).split()))
    b_tokens = " ".join(sorted(normalise(b).split()))
    return SequenceMatcher(None, a_tokens, b_tokens).ratio() * 100

def is_match(expected, actual, threshold=85, is_date=False):
    """Return True if values match (exact or fuzzy above threshold)."""
    if is_date:
        return normalise_date(expected) == normalise_date(actual)
    e = normalise(expected)
    a = normalise(actual)
    if not e and not a:
        return True   # both empty → match
    if not e or not a:
        return False  # one empty, one not → miss
    if e == a:
        return True
    return fuzzy_score(e, a) >= threshold

def fmt_pct(n, d):
    return f"{n/d*100:.1f}%" if d > 0 else "N/A"

# ---------------------------------------------------------------------------
# Layer B: Parse Word document directly
# ---------------------------------------------------------------------------

def parse_docx_ground_truth(docx_path):
    """
    Extract critical structured fields from each patient table in the docx.
    Returns a list of dicts (one per patient, in document order).
    """
    doc = docx.Document(docx_path)
    patients = []

    for table in doc.tables:
        # Row 0: headers  |  Row 1: patient details cell
        if len(table.rows) < 2:
            continue
        header_text = " ".join(c.text for c in table.rows[0].cells).lower()
        if 'patient details' not in header_text:
            continue

        details_cell = table.rows[1].cells[0].text.strip()

        # --- MRN ---
        mrn = ""
        m = re.search(r'Hospital Number:\s*([^\n]+)', details_cell)
        if m:
            mrn = re.sub(r'\([a-zA-Z]\)', '', m.group(1)).strip()

        # --- NHS Number ---
        nhs = ""
        m = re.search(r'NHS Number:\s*([^\n]+)', details_cell)
        if m:
            nhs = re.sub(r'\([a-zA-Z]\)', '', m.group(1)).strip()
            # Guard: if it picked up the next label line, discard
            if re.match(r'^[A-Za-z]', nhs) and not re.match(r'^\d', nhs):
                nhs = ""

        # --- Name (two formats) ---
        name = ""
        # Format A: "Name: Firstname Lastname"
        m = re.search(r'\bName:\s*([^\n]+)', details_cell)
        if m:
            name = re.sub(r'\([a-zA-Z]\)', '', m.group(1)).strip()
        else:
            # Format B: ALL CAPS line not starting with known labels
            lines = [l.strip() for l in details_cell.split('\n') if l.strip()]
            for line in lines:
                if (re.match(r'^[A-Z][A-Z\s\'\-]+$', line) and
                        not any(line.startswith(x) for x in
                                ['HOSPITAL', 'NHS', 'MALE', 'FEMALE'])):
                    name = re.sub(r'\([a-zA-Z]\)', '', line).strip()
                    break

        # --- Gender (two formats) ---
        gender = ""
        m = re.search(r'Gender:\s*(\w+)', details_cell)
        if m:
            gender = m.group(1).capitalize()
        else:
            m = re.search(r'\b(Male|Female)\b', details_cell, re.IGNORECASE)
            if m:
                gender = m.group(1).capitalize()

        # --- DOB (two formats) ---
        dob = ""
        # Format A: "DOB: DD/MM/YYYY"
        m = re.search(r'DOB:\s*(\d{2}/\d{2}/\d{4})', details_cell)
        if m:
            dob = m.group(1)
        else:
            # Format B: bare date at start of line e.g. "26/05/1970(a) Age: 55"
            m = re.search(r'(\d{2}/\d{2}/\d{4})', details_cell)
            if m:
                dob = m.group(1)

        # --- MDT Meeting Date (from paragraph just before table) ---
        mdt_date = ""
        # Find MDT date from paragraph text in the document
        # It appears as "Colorectal Multidisciplinary Meeting DD/MM/YYYY(i)"
        for para in doc.paragraphs:
            m = re.search(r'Multidisciplinary Meeting\s+(\d{2}/\d{2}/\d{4})', para.text)
            if m:
                mdt_date = m.group(1)
                # We'll assign per-table below after we know table order

        # --- Diagnosis (from row 3, col 0) ---
        diagnosis = ""
        if len(table.rows) >= 4:
            diag_cell = table.rows[3].cells[0].text
            m = re.search(r'Diagnosis:\s*(.+?)(?:\n|ICD10|$)', diag_cell)
            if m:
                val = m.group(1).strip()
                # Skip if empty or just whitespace
                if val and not re.match(r'^\s*$', val):
                    diagnosis = val

        patients.append({
            'mrn': mrn,
            'nhs': nhs,
            'name': name,
            'gender': gender,
            'dob': dob,
            'diagnosis': diagnosis,
        })

    # --- Assign MDT dates properly by matching paragraphs to table order ---
    # Re-parse paragraphs in document order; each meeting heading precedes its tables
    current_mdt_date = ""
    patient_idx = 0
    mdt_dates_ordered = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            text = element.text_content() if hasattr(element, 'text_content') else ''
            # Use lxml text extraction
            text = ''.join(element.itertext())
            m = re.search(r'Multidisciplinary Meeting\s+(\d{2}/\d{2}/\d{4})', text)
            if m:
                current_mdt_date = m.group(1)
        elif tag == 'tbl':
            mdt_dates_ordered.append(current_mdt_date)

    for i, p in enumerate(patients):
        if i < len(mdt_dates_ordered):
            p['mdt_date'] = mdt_dates_ordered[i]
        else:
            p['mdt_date'] = ""

    return patients


# ---------------------------------------------------------------------------
# Layer A: JSON-based validation
# ---------------------------------------------------------------------------

# Map: JSON field path → Excel column name
# path uses dot notation, e.g. "patient_details.dob"
JSON_TO_EXCEL_MAP = {
    'patient_details.dob':             'Demographics: \nDOB(a)',
    'patient_details.hospital_number': 'Demographics: MRN(c)',
    'patient_details.nhs_number':      'Demographics: \nNHS number(d)',
    'patient_details.gender':          'Demographics: \nGender(e)',
    'mdt_meeting_date':                '1st MDT: date(i)',
    'staging_and_diagnosis.diagnosis': 'Histology: Biopsy result(g)',
    'clinical_details':                'Endoscopy: Findings(f)',
    'mdt_outcome':                     '1st MDT: Treatment approach \n(TNT, downstaging chemotherapy, downstaging nCRT, downstaging shortcourse RT, Papillon +/- EBRT, straight to surgery(h)',
}

def get_nested(data, path):
    """Retrieve value from nested dict using dot-notation path."""
    parts = path.split('.')
    val = data
    for p in parts:
        if isinstance(val, dict):
            val = val.get(p)
        else:
            return None
    return val


def run_layer_a(json_dir, df_excel):
    """Compare JSON fields to Excel output for all 50 cases."""
    json_files = sorted(Path(json_dir).glob('case_*.json'))
    if not json_files:
        print(f"  [WARNING] No JSON files found in {json_dir}")
        return None

    results = []
    field_scores = {path: {'match': 0, 'miss': 0, 'empty': 0} for path in JSON_TO_EXCEL_MAP}

    for i, jf in enumerate(json_files):
        with open(jf) as f:
            jdata = json.load(f)

        if i >= len(df_excel):
            break

        excel_row = df_excel.iloc[i]
        row_result = {'case': jf.stem, 'fields': {}}

        for json_path, excel_col in JSON_TO_EXCEL_MAP.items():
            json_val = get_nested(jdata, json_path)
            excel_val = excel_row.get(excel_col, None)

            j_norm = normalise(json_val)
            e_norm = normalise(excel_val)

            if not j_norm and not e_norm:
                status = 'both_empty'
                field_scores[json_path]['empty'] += 1
            elif not e_norm:
                status = 'excel_empty'
                field_scores[json_path]['miss'] += 1
            elif not j_norm:
                status = 'json_empty'
                field_scores[json_path]['empty'] += 1
            elif is_match(j_norm, e_norm):
                status = 'match'
                field_scores[json_path]['match'] += 1
            else:
                status = 'mismatch'
                field_scores[json_path]['miss'] += 1
                score = fuzzy_score(j_norm, e_norm)
                row_result['fields'][json_path] = {
                    'status': status,
                    'score': round(score, 1),
                    'json': str(json_val)[:80],
                    'excel': str(excel_val)[:80]
                }
                continue

            row_result['fields'][json_path] = {'status': status}

        results.append(row_result)

    return results, field_scores


# ---------------------------------------------------------------------------
# Layer B: Compare docx ground truth to Excel
# ---------------------------------------------------------------------------

DOCX_TO_EXCEL_MAP = {
    'mrn':       'Demographics: MRN(c)',
    'nhs':       'Demographics: \nNHS number(d)',
    'gender':    'Demographics: \nGender(e)',
    'dob':       'Demographics: \nDOB(a)',
    'mdt_date':  '1st MDT: date(i)',
    'diagnosis': 'Histology: Biopsy result(g)',
}

def run_layer_b(docx_path, df_excel):
    """Independent docx parsing vs Excel for critical fields."""
    gt_patients = parse_docx_ground_truth(docx_path)

    field_scores = {f: {'match': 0, 'miss': 0, 'empty': 0} for f in DOCX_TO_EXCEL_MAP}
    mismatches = []

    for i, gt in enumerate(gt_patients):
        if i >= len(df_excel):
            break
        excel_row = df_excel.iloc[i]

        for field, excel_col in DOCX_TO_EXCEL_MAP.items():
            gt_val = gt.get(field, '')
            excel_val = excel_row.get(excel_col, None)

            g_norm = normalise(gt_val)
            e_norm = normalise(excel_val)
            is_date_field = field in ('dob', 'mdt_date')

            if not g_norm and not e_norm:
                field_scores[field]['empty'] += 1
            elif not e_norm:
                field_scores[field]['miss'] += 1
                mismatches.append({
                    'case': f'case_{i:03d}',
                    'field': field,
                    'ground_truth': gt_val,
                    'excel': str(excel_val),
                    'issue': 'excel_empty'
                })
            elif is_match(gt_val, excel_val, is_date=is_date_field):
                field_scores[field]['match'] += 1
            else:
                field_scores[field]['miss'] += 1
                mismatches.append({
                    'case': f'case_{i:03d}',
                    'field': field,
                    'ground_truth': gt_val,
                    'excel': str(excel_val),
                    'issue': 'mismatch',
                    'score': round(fuzzy_score(g_norm, e_norm), 1)
                })

    return field_scores, mismatches


# ---------------------------------------------------------------------------
# Report generation
# ---------------------------------------------------------------------------

def print_report(layer_a_results, layer_a_field_scores,
                 layer_b_field_scores, layer_b_mismatches,
                 n_patients):

    sep = "=" * 65

    print(f"\n{sep}")
    print("  V5 PIPELINE VALIDATION REPORT — STEP 1")
    print(sep)
    print(f"  Patients evaluated : {n_patients}")
    print(f"  Validation date    : {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")

    # --- Layer A ---
    print(f"\n{sep}")
    print("  LAYER A — JSON-Based Broad Coverage Validation")
    print(sep)
    if layer_a_field_scores:
        total_match = total_miss = total_empty = 0
        print(f"\n  {'Field':<55} {'Match':>6} {'Miss':>6} {'Empty':>6} {'Acc':>7}")
        print(f"  {'-'*55} {'-'*6} {'-'*6} {'-'*6} {'-'*7}")
        for path, s in layer_a_field_scores.items():
            total = s['match'] + s['miss']
            acc = fmt_pct(s['match'], total) if total > 0 else "—"
            label = path.split('.')[-1][:54]
            print(f"  {label:<55} {s['match']:>6} {s['miss']:>6} {s['empty']:>6} {acc:>7}")
            total_match += s['match']
            total_miss  += s['miss']
            total_empty += s['empty']

        total = total_match + total_miss
        layer_a_acc = total_match / total * 100 if total > 0 else 0
        print(f"\n  {'TOTAL':<55} {total_match:>6} {total_miss:>6} {total_empty:>6} {fmt_pct(total_match, total):>7}")
        print(f"\n  Layer A Accuracy: {layer_a_acc:.1f}%")
    else:
        print("  [SKIPPED] No JSON files provided.")
        layer_a_acc = None

    # --- Layer B ---
    print(f"\n{sep}")
    print("  LAYER B — Independent Docx Ground Truth Audit")
    print(sep)
    total_match = total_miss = total_empty = 0
    print(f"\n  {'Field':<20} {'Match':>6} {'Miss':>6} {'Empty':>6} {'Acc':>7}")
    print(f"  {'-'*20} {'-'*6} {'-'*6} {'-'*6} {'-'*7}")
    for field, s in layer_b_field_scores.items():
        total = s['match'] + s['miss']
        acc = fmt_pct(s['match'], total) if total > 0 else "—"
        print(f"  {field:<20} {s['match']:>6} {s['miss']:>6} {s['empty']:>6} {acc:>7}")
        total_match += s['match']
        total_miss  += s['miss']
        total_empty += s['empty']

    total = total_match + total_miss
    layer_b_acc = total_match / total * 100 if total > 0 else 0
    print(f"\n  {'TOTAL':<20} {total_match:>6} {total_miss:>6} {total_empty:>6} {fmt_pct(total_match, total):>7}")
    print(f"\n  Layer B Accuracy: {layer_b_acc:.1f}%")

    if layer_b_mismatches:
        print(f"\n  Mismatches / Missing values ({len(layer_b_mismatches)} total):")
        for m in layer_b_mismatches[:15]:
            score_str = f"  [fuzzy:{m.get('score','—')}]" if 'score' in m else ""
            print(f"    {m['case']} | {m['field']:<12} | GT: {str(m['ground_truth'])[:30]:<30} | Excel: {str(m['excel'])[:30]}{score_str}")
        if len(layer_b_mismatches) > 15:
            print(f"    ... and {len(layer_b_mismatches)-15} more (see full JSON report)")

    # --- Combined Score ---
    print(f"\n{sep}")
    print("  COMBINED SCORE")
    print(sep)
    if layer_a_acc is not None:
        combined = layer_a_acc * 0.5 + layer_b_acc * 0.5
        print(f"\n  Layer A (50% weight): {layer_a_acc:.1f}%")
        print(f"  Layer B (50% weight): {layer_b_acc:.1f}%")
        print(f"\n  ★  OVERALL STEP 1 SCORE: {combined:.1f}%")
    else:
        print(f"\n  Layer B only: {layer_b_acc:.1f}%")
        print(f"\n  ★  OVERALL STEP 1 SCORE (Layer B only): {layer_b_acc:.1f}%")

    print(f"\n{sep}\n")
    return layer_b_acc, layer_a_acc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description='V5 Pipeline Validation — Step 1')
    parser.add_argument('--docx',  default='data/hackathon-mdt-outcome-proformas.docx')
    parser.add_argument('--excel', default='v5_VLM_LLM_DirectLogic/output/hackathon_output_final.xlsx')
    parser.add_argument('--jsons', default='v5_VLM_LLM_DirectLogic/output/json_raw_claude/',
                        help='Directory of case_000.json ... case_049.json')
    parser.add_argument('--out',   default='validation_report_step1.json',
                        help='Path to save full JSON report')
    args = parser.parse_args()

    print(f"\nLoading Excel output: {args.excel}")
    df = pd.read_excel(args.excel)
    n = len(df)
    print(f"  → {n} rows, {len(df.columns)} columns")

    # Layer A
    print(f"\nRunning Layer A (JSON validation) from: {args.jsons}")
    if os.path.isdir(args.jsons):
        la_results, la_field_scores = run_layer_a(args.jsons, df)
        print(f"  → Evaluated {len(la_results)} cases")
    else:
        print(f"  [SKIP] JSON directory not found: {args.jsons}")
        la_results, la_field_scores = None, None

    # Layer B
    print(f"\nRunning Layer B (Docx ground truth audit) from: {args.docx}")
    lb_field_scores, lb_mismatches = run_layer_b(args.docx, df)
    print(f"  → Parsed {n} patients from docx")

    # Report
    lb_acc, la_acc = print_report(
        la_results, la_field_scores,
        lb_field_scores, lb_mismatches,
        n_patients=n
    )

    # Save full JSON report
    report = {
        'layer_a': {
            'field_scores': la_field_scores,
            'accuracy_pct': round(la_acc, 2) if la_acc is not None else None,
        },
        'layer_b': {
            'field_scores': lb_field_scores,
            'mismatches': lb_mismatches,
            'accuracy_pct': round(lb_acc, 2),
        },
        'combined_score_pct': round(
            (la_acc * 0.5 + lb_acc * 0.5) if la_acc is not None else lb_acc, 2
        )
    }
    with open(args.out, 'w') as f:
        json.dump(report, f, indent=2)
    print(f"Full JSON report saved to: {args.out}\n")


if __name__ == '__main__':
    main()
